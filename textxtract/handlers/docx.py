"""DOCX file handler for comprehensive text extraction.

This handler extracts text from:
- Document paragraphs
- Tables and cells
- Headers and footers
- Text boxes and shapes
- Footnotes and endnotes (if available)
"""

from pathlib import Path
from typing import Optional
import re

from textxtract.core.base import FileTypeHandler
from textxtract.core.exceptions import ExtractionError


class DOCXHandler(FileTypeHandler):
    """Enhanced handler for comprehensive text extraction from DOCX files.
    
    This handler provides complete text extraction from Microsoft Word documents,
    including all document elements such as paragraphs, tables, headers, footers,
    text boxes, and footnotes. It's designed to handle complex document layouts
    commonly found in resumes, reports, and structured documents.
    
    Features:
        - Extracts text from document body paragraphs
        - Processes table content with cell-by-cell extraction
        - Captures header and footer text from all sections
        - Attempts to extract text from embedded text boxes and shapes
        - Handles footnotes and endnotes when available
        - Deduplicates repeated content
        - Cleans and normalizes extracted text
    
    Example:
        >>> handler = DOCXHandler()
        >>> text = handler.extract(Path("document.docx"))
        >>> print(text)
        "Document title\nParagraph content...\nTable data | Column 2..."
        
        >>> # Async extraction
        >>> text = await handler.extract_async(Path("document.docx"))
    """

    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text.
        
        Performs various text cleaning operations to improve readability
        and consistency of extracted content.
        
        Args:
            text (str): Raw text to be cleaned.
            
        Returns:
            str: Cleaned and normalized text with proper spacing and formatting.
            
        Note:
            - Normalizes multiple whitespace characters to single spaces
            - Removes excessive consecutive dots/periods
            - Fixes spacing around punctuation marks
            - Strips leading and trailing whitespace
        """
        if not text:
            return ""
        
        # Normalize whitespace (replace multiple spaces, tabs, newlines with single space)
        text = re.sub(r'\s+', ' ', text)
        # Remove excessive dots/periods (likely formatting artifacts)
        text = re.sub(r'\.{2,}', ' ', text)
        # Clean up spacing around punctuation (remove spaces before punctuation)
        text = re.sub(r'\s+([.!?,:;])', r'\1', text)
        return text.strip()

    def extract(self, file_path: Path, config: Optional[dict] = None) -> str:
        """Extract text from a DOCX file with comprehensive content capture.
        
        Performs thorough text extraction from all available document elements
        including body text, tables, headers, footers, and embedded content.
        
        Args:
            file_path (Path): Path to the DOCX file to extract text from.
            config (Optional[dict], optional): Configuration options for extraction.
                Currently not used but reserved for future enhancements.
                
        Returns:
            str: Extracted and cleaned text from the document with proper formatting.
                Returns empty string if no text is found.
                
        Raises:
            ExtractionError: If the file cannot be read or processed, or if the
                python-docx library is not available.
                
        Note:
            - Text is deduplicated to avoid repeated content from overlapping elements
            - Table content is formatted with pipe separators between columns
            - Special content (footnotes, text boxes) is labeled with descriptive tags
            - Sentence breaks are automatically inserted for better readability
        """
        try:
            from docx import Document
            import re

            # Load the document
            doc = Document(file_path)
            text_parts = []
            processed_text = set()  # Track processed text to avoid duplicates
            
            # Extract text from main document paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text and text not in processed_text:
                    text_parts.append(text)
                    processed_text.add(text)
            
            # Extract text from all tables in the document
            for table in doc.tables:
                table_texts = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        # Process each paragraph within the cell
                        cell_paragraphs = []
                        for paragraph in cell.paragraphs:
                            text = paragraph.text.strip()
                            if text and text not in processed_text:
                                cell_paragraphs.append(text)
                                processed_text.add(text)
                        if cell_paragraphs:
                            row_text.append(" ".join(cell_paragraphs))
                    if row_text:
                        # Join cell contents with pipe separator for table structure
                        table_texts.append(" | ".join(row_text))
                
                # Add table content to main text collection
                if table_texts:
                    text_parts.extend(table_texts)
            
            # Extract text from headers and footers across all document sections
            for section in doc.sections:
                # Process header content
                if section.header:
                    for paragraph in section.header.paragraphs:
                        text = paragraph.text.strip()
                        if text and text not in processed_text:
                            text_parts.append(text)
                            processed_text.add(text)
                
                # Process footer content
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        text = paragraph.text.strip()
                        if text and text not in processed_text:
                            text_parts.append(text)
                            processed_text.add(text)
            
            # Attempt to extract footnotes and endnotes (may not be available in all documents)
            try:
                # Extract footnotes if present
                if hasattr(doc, 'footnotes'):
                    for footnote in doc.footnotes:
                        for paragraph in footnote.paragraphs:
                            text = paragraph.text.strip()
                            if text and text not in processed_text:
                                text_parts.append(f"[Footnote: {text}]")
                                processed_text.add(text)
                
                # Extract endnotes if present
                if hasattr(doc, 'endnotes'):
                    for endnote in doc.endnotes:
                        for paragraph in endnote.paragraphs:
                            text = paragraph.text.strip()
                            if text and text not in processed_text:
                                text_parts.append(f"[Endnote: {text}]")
                                processed_text.add(text)
            except Exception:
                # Footnote/endnote extraction is optional - continue if it fails
                pass
            
            # Attempt to extract text from embedded text boxes and shapes using XML parsing
            try:
                from docx.oxml.ns import qn
                
                # Iterate through document XML elements to find drawing content
                for element in doc.element.body.iter():
                    if element.tag.endswith('}txbxContent'):
                        # Extract text from text box elements
                        for para in element.iter():
                            if para.tag.endswith('}t') and para.text:
                                text = para.text.strip()
                                if text and text not in processed_text:
                                    text_parts.append(f"[TextBox: {text}]")
                                    processed_text.add(text)
            except Exception:
                # Text box extraction is optional - continue if it fails
                pass
            
            # Process and format the final output
            if text_parts:
                # Clean each text part and filter out empty content
                cleaned_parts = [self._clean_text(part) for part in text_parts if part.strip()]
                result = "\n".join(cleaned_parts)
                
                # Add proper sentence breaks for improved readability
                result = re.sub(r'([.!?])\s*([A-Z])', r'\1\n\2', result)
                return result.strip()
            
            return ""
            
        except Exception as e:
            raise ExtractionError(f"DOCX extraction failed: {e}")

    async def extract_async(
        self, file_path: Path, config: Optional[dict] = None
    ) -> str:
        """Asynchronously extract text from a DOCX file.
        
        Provides non-blocking text extraction by running the synchronous
        extraction method in a separate thread.
        
        Args:
            file_path (Path): Path to the DOCX file to extract text from.
            config (Optional[dict], optional): Configuration options for extraction.
                Currently not used but reserved for future enhancements.
                
        Returns:
            str: Extracted and cleaned text from the document with proper formatting.
                Returns empty string if no text is found.
                
        Raises:
            ExtractionError: If the file cannot be read or processed, or if the
                python-docx library is not available.
                
        Note:
            This method uses asyncio.to_thread() to run the synchronous extraction
            in a thread pool, making it suitable for async/await usage patterns.
        """
        import asyncio

        return await asyncio.to_thread(self.extract, file_path, config)
